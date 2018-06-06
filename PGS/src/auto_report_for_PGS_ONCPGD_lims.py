#!/data/software/bin/python3
# -*- coding: utf-8 -*-
# 重构自：/data/software/reporter/PGS_report_v2.0_for_LIMS_latest.py
__author__ = 'zhangdawei@yikongenomics.com'
#__version__ = "v3.0 || 2017/12/21 + 2017/12/22 + 2017/12/26 + 2017/12/28"
#__version__ = "v3.1 || 2018/1/2" # 增加了IBPGS ONPGS CPGD
#__version__ = 'v3.2 || 2018/1/12' # 使送检单中没有的样本也能出现在报告中
#__version__ = 'old || 2018/1/19' #  使用旧模板，使用旧的四列的分析结果
#__version__ = 'old_v1.2 || 2018/2/5' #  cnv结果变为三列
#__version__ = 'old_v1.3 || 2018/3/2' # 由送检barcode和结果barcode，生成一个新的用于出结果的barcode
#__version__ = 'old_v1.4 || 2018/3/31' # out.txt变成了21列，最好能同时支持20列和21列
__version__ = 'old_v1.5 || 2018/4/28' # 由默认带logo改成默认不带logo
'''
requirements:
	python 2.7
	pip install python-docx
	pip install openpyxl
	pip install docxtpl
'''
import sys
import os
import errno
import os.path
import argparse
import datetime
import subprocess
import shutil
# used for create .docx file
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
# used for vert table alignment
from docx.oxml.shared import OxmlElement, qn
# used for create excel file 
from openpyxl import Workbook
from docxtpl import DocxTemplate
import logging

# 报告模板
#PGS_REPORT_TEMPLATE = "/data/software/reporter/template/"

PGS_REPORT_TEMPLATE = "/data/software/Pipeline/auto_report_pipeline/REPORT_TEMPLATE_OLD/with_signature/PGS_report_template_with_signature_v170719.2.docx"
PGS_REPORT_TEMPLATE_WITHOUT_LOGO = "/data/software/Pipeline/auto_report_pipeline/REPORT_TEMPLATE_OLD/PGS_report_template_without_logo.docx"
ONCPGD_REPORT_TEMPLATE = "/data/software/Pipeline/auto_report_pipeline/REPORT_TEMPLATE_OLD/with_signature/ONCPGD_report_template_with_signature_v170405.5.docx"
CPGD_REPORT_TEMPLATE = "/data/software/Pipeline/auto_report_pipeline/REPORT_TEMPLATE_OLD/with_signature/CPGD_report_template_with_signature_v170719.2.docx"
# lims送检单缺省值
MISSING_VALUES = [u"无"]

# 医院定制化：不写医院名的
HOSPITAL_ID_WITHOUT_HOSPITAL_NAME = ['07A']

# 写log
def create_logging(log_name):
        # 创建一个logger
        logger = logging.getLogger(log_name)
        logger.setLevel(logging.DEBUG)
        ch = logging.StreamHandler()
        ch.setLevel(logging.DEBUG)
        # 定义handler的输出格式
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        ch.setFormatter(formatter)
        # 给logger添加handler
        logger.addHandler(ch)
        return logger

# 格式：调整excel列宽
def adjustColumnWidth(ws):
	dims = {}
	for row in ws.rows:
		for cell in row:
			if cell.value:
				dims[cell.column] = max((dims.get(cell.column, 0), len(cell.value)))
	for col, value in dims.items():
		ws.column_dimensions[col].width = value+5
	return ws

# 格式：word中的时间
def format_time(time):
	if time == "":
		return time
	else:
		tmp = time.split('-')
		sample_submit_date = u"{}年{}月{}日".format(tmp[0],tmp[1],tmp[2])
		return(sample_submit_date)

# 信息抓取1： 
# 解析20列的out.txt, 生成了dict_patient_info,dict_sample_info,sample_barcodes
# dict_patient_info 哈希，抓取了所有 20列 信息
# dict_sample_info 哈希， barcode -> 样本名
# sample_barcodes 数组，所有barcode
### 2018/3/31 由20列变为21列


def parse_patient_info(info_file):
	dict_patient_info={}
	dict_sample_info={}
	sample_ids = []
	sample_barcodes = []
	sample_types = []
	fh_patient_info = open(info_file)
	next(fh_patient_info)
	for line in fh_patient_info.readlines():
		if line.split():
			#line=line.strip()
			tmp_list=line.split('\t')
			#if len(tmp_list) == 20:
			if len(tmp_list) >= 20:
				dict_patient_info["SampleSheetBarcode"]= tmp_list[0].strip()
				dict_patient_info["ProjectID"]= tmp_list[1].strip() ###项目名
				dict_patient_info["WomanName"]= tmp_list[2].strip()
				dict_patient_info["ManName"]= tmp_list[3].strip()
				dict_patient_info["WomanAge"]= tmp_list[4].strip()
				dict_patient_info["ManAge"]= tmp_list[5].strip()
				dict_patient_info["SubmissionOrganization"]= tmp_list[6].strip()
				dict_patient_info["Doctor"] = tmp_list[7].strip()
				dict_patient_info["BiopsyDate"]= format_time(tmp_list[8].strip())
				dict_patient_info["SubmissionDate"]= format_time(tmp_list[9].strip())
				dict_patient_info["TestType"]= tmp_list[10].strip()
				sample_types.append(tmp_list[11].strip()) #样本类型
				dict_patient_info["ManKaryotypeInfo"]= tmp_list[12].strip()
				dict_patient_info["ManKaryotype"]= tmp_list[13].strip()
				dict_patient_info["WomanKaryotypeInfo"]= tmp_list[14].strip()
				dict_patient_info["WomanKaryotype"]= tmp_list[15].strip()

				if dict_patient_info["ManKaryotype"] == "":
					dict_patient_info["ManKaryotype"] = "无"
				if dict_patient_info["WomanKaryotype"] == "":
					dict_patient_info["WomanKaryotype"] = "无"
				dict_patient_info["Karyotype"] = "男方：" + dict_patient_info["ManKaryotype"] + "; " + "女方：" + dict_patient_info['WomanKaryotype']
				
				#样本名，如果送检单里没有写，就取barcode替代
				if tmp_list[16].strip():
					sample_ids.append(tmp_list[16].strip()) 
				else:
					sample_ids.append(tmp_list[17].strip()) 

				sample_barcodes.append(tmp_list[17].strip()) ###样本barcode
				dict_patient_info["AnalysisType"]= tmp_list[18].strip()
				dict_patient_info["Template"] = tmp_list[19].strip()
				
			else:
#				print("sample info's No. columns are not equal 20!")
				print("sample info's No. columns are not >= 20!")
				exit()
		else:
			pass
	dict_patient_info["SampleType"] = ','.join(set(sample_types)) # tmp_list[11]样本类型
	dict_patient_info["sample_ids"] = sample_ids # tmp_list[16]样本名
	dict_patient_info["sample_barcodes"] = sample_barcodes # tmp_list[17]样本barcode
	
	dict_sample_info = dict(zip(sample_barcodes, sample_ids))
	return dict_patient_info,dict_sample_info,sample_barcodes

# 信息抓取2：
# 抓取cnv.txt
# result_barcodes [barcode]
# dict_result {barcode: 核型信息}
# dict_result_gender {barcode: 性别核型}

def parse_cnv_file(cnv_file, qc):
	result_barcodes = []
	dict_result = {}
	dict_result_gender = {}
	
	fh2 = open(cnv_file)
	for line in fh2.readlines():
		if line.split():
			line=line.rstrip()
			tmp_list = line.split('\t')
			result_barcodes.append(tmp_list[0].strip())
			dict_result[tmp_list[0]] = tmp_list[1] # 核型
			dict_result_gender[tmp_list[0]] = tmp_list[2] # 性别核型
			if 'FAIL' in qc[tmp_list[0]]:
				dict_result[tmp_list[0]] = 'N/A'  # 核型
				dict_result_gender[tmp_list[0]] = 'N/A'  # 性别核型
		else:
			pass
	fh2.close()
	return (result_barcodes,dict_result,dict_result_gender)

# 信息抓取3：
# cnv图抓取
def parse_graph(project_dir,graph_tag,xy_tag,result_barcodes,bin_size):
	dict_png = {}
	for barcode in result_barcodes:
		png = project_dir + '/analysis/' + graph_tag + '/' + xy_tag + '/' + barcode + '_' + bin_size + '_' + xy_tag + '.png'
		dict_png[barcode] = png
	return dict_png

# 2018/3/2补充，由送检barcode和结果barcode，生成一个新的用于出结果的barcode
def gen_report_barcodes(sample_barcodes, result_barcodes):
	report_barcodes = []
	for barcode in sample_barcodes:
		if barcode in result_barcodes:
			report_barcodes.append(barcode)
		else:
			pass
	
	for barcode in result_barcodes:
		if barcode not in sample_barcodes:
			report_barcodes.append(barcode)
		else:
			pass

	return report_barcodes

def main():
	
### MARK1 抓取命令行参数	
	parser = argparse.ArgumentParser(prog = 'PGS_report', description = 'Save PGS project\'s graph (.png) file to a (.docx) word file and sex info file to an excel (.xlsx) file and write a report based on word template to a (.docx) file. ')
	parser.add_argument('--project_type', action = "store",required=False,default="PGS",choices=["PGS","ChromInst","ONCPGD","ONPGS","CPGD","IBPGS"], help = "The Project type. [default=PGS]")
	parser.add_argument('--project_dir', action = "store",required=True, help = 'The project output directory')
	parser.add_argument('--bin_size', action = "store",required=True, help = 'The  bin size, like 1000K. ')
	parser.add_argument('--cnv_file', action = "store",required=True, help = 'The call cnv file created by pipeline. It is used to create the sex info excel file and PGS test results are extracted from it.')
	parser.add_argument('--patient_info',action = "store",required=True, help = 'The patient info txt file')
	parser.add_argument('--out_dir', action = "store",required=True, help = 'The output graph/info/report files\'s directory ')
	parser.add_argument('--data_sts', action="store", required=True, help="data.sts file")
	args = parser.parse_args()
	project_type = args.project_type
	project_dir = os.path.abspath(args.project_dir) 
	bin_size = args.bin_size
	cnv_file = os.path.abspath(args.cnv_file)
	patient_info = os.path.abspath(args.patient_info)
	out_dir = os.path.abspath(args.out_dir)
	data_sts = os.path.abspath(args.data_sts)

### MARK2 解析参数文件，转化为变量
	logger = create_logging('PGS report logger')

##  MARK2.1 解析送检单 patient_info
	dict_patient_info,dict_sample_info,sample_barcodes = parse_patient_info(patient_info)
	logger.info("DONE:parse patient info")
#   MARK2.1.1 将缺失值“单击此处输入文字”替换为空字符
	for (k,v) in dict_patient_info.items():
		if v in MISSING_VALUES:
			dict_patient_info[k] = ''
	logger.info("DONE:clean patient info")

# 信息抓取补充
######read qc_discription.txt
	dict_qc = {}
	fh_sts = open(data_sts)
	for line in fh_sts.readlines():
		if line.split():
			line = line.rstrip()
			list_sts = line.split("\t")
			dict_qc[list_sts[0]] = list_sts[10]
	logger.info("DONE:parse data.sts")

##  MARK2.2 解析结果 cnv_file
	result_barcodes,dict_result,dict_result_gender = parse_cnv_file(cnv_file, dict_qc)
	logger.info("DONE:parse cnv file")

##  MARK2.3 解析cnv图
	dict_red_blue_png_no_XY = parse_graph(project_dir,'graph','with_chrID_no_XY',result_barcodes,bin_size)
	dict_red_blue_png_XY = parse_graph(project_dir,'graph','with_chrID_with_XY',result_barcodes,bin_size)
	dict_colorful_png_XY = parse_graph(project_dir,'graph1','with_chrID_with_XY',result_barcodes,bin_size)
	logger.info("DONE:parse png graph")

### MARK3 根据需求，生成内部参量
##  MARK 2018/3/1补充，由送检barcode和结果barcode，生成一个新的用于出结果的barcode
	report_barcodes = gen_report_barcodes(sample_barcodes, result_barcodes)


##  MARK3.1 抓取 项目名，医院名，女方姓名，是否要logo
#   sample_sheet_ID 项目名	
	sample_sheet_ID = dict_patient_info["ProjectID"]
	if sample_sheet_ID in MISSING_VALUES:
		logger.error("Project ID is missing!")
		exit(1)
#   hospital_ID 医院名
	array_sample_sheet_ID = sample_sheet_ID.split('_')
	hospital_ID = array_sample_sheet_ID[3]
#   woman_name 女方姓名
	woman_name = dict_patient_info['WomanName']
	if 'Control' in sample_sheet_ID or 'control' in sample_sheet_ID:
			woman_name = 'Control'
#   if_logo 是否要logo
	if_logo = dict_patient_info['Template']

##  MARK3.2 完整的输出名
	out_name = "Project_" + sample_sheet_ID + u"（" + woman_name + u"）"
	if woman_name == '':
		out_name = "Project_" + sample_sheet_ID

	
### MARK4 结果展示

#######################
##	MARK4.1 生成CNV全图
#   CNV全图名
	out_graph = out_dir + '/' + out_name + u"CNV全图.docx"

	document = Document()
	for barcode in report_barcodes:
		png_fullpath = dict_colorful_png_XY[barcode];
		if os.path.exists(png_fullpath):
			document.add_picture(png_fullpath,width=Inches(6.35))
		else:
			log = "save graph file : "+png_fullpath+" does not exist!"
			logger.warning(log)
	document.save(out_graph)
	logger.info("DONE:save graph file")

############################################################
##  MARK4.2 生成性别信息表
#   性别信息名
	out_info = out_dir + '/' + out_name + "info.xlsx"

	wb = Workbook()
	ws = wb.active
	ws.title = "info"
	excel_row=1
	for barcode in report_barcodes:
		ws.cell(row=excel_row, column=1, value=barcode)
		ws.cell(row=excel_row, column=2, value=dict_sample_info.get(barcode, barcode))
		ws.cell(row=excel_row, column=3, value=dict_result_gender[barcode])
		excel_row+=1
	ws = adjustColumnWidth(ws)
	wb.save(out_info)
	logger.info("DONE:save sex info xlsx file")

####################
##  MARK4.3 生成报告

##  MARK 报告名
	project_type_name = ''
	
	if project_type == 'ONCPGD' or project_type == 'ONPGS':
		project_type_name = "24h-胚胎染色体拷贝数检测报告单"
	elif project_type == 'CPGD':
		project_type_name = "MALBAC-PGD™ 染色体病胚胎植入前遗传学诊断报告单"
	elif project_type == 'PGS' or project_type == 'IBPGS' :
		project_type_name = "胚胎植入前遗传学筛查（PGS）检测报告单"
	else:
		project_type_name = "ChromInst 9h-胚胎染色体拷贝数检测报告单"
	out_report = out_dir + '/' + out_name + project_type_name + ".docx"

##  MARK 报告模板
	if project_type == 'ONCPGD' or project_type == 'ONPGS':	
		report_temp  = ONCPGD_REPORT_TEMPLATE
	elif project_type == 'CPGD':
		report_temp  = CPGD_REPORT_TEMPLATE
	else:
		if if_logo == "Yes" or if_logo == "yes":
			report_temp  = PGS_REPORT_TEMPLATE
		else:
			report_temp  = PGS_REPORT_TEMPLATE_WITHOUT_LOGO

#   MARK4.3.1 定制化
#   某些医院要求不写医院名
	if hospital_ID in HOSPITAL_ID_WITHOUT_HOSPITAL_NAME:
		dict_patient_info["SubmissionOrganization"] = ''
#   ONCPGD活检日期填到送检日期位置
	if project_type == 'ONCPGD' or project_type == 'ONPGS':
		dict_patient_info["SubmissionDate"] = dict_patient_info["BiopsyDate"]

### MARK4.3.4 取模版，贴报告
	shutil.copyfile(report_temp , out_report)
	ReportTML = DocxTemplate(out_report)

##  报_送检信息
#   获取报告日期
	now = datetime.datetime.now()
	month = now.month
	day = now.day
	if month < 10:
		month = '0' + str(month)
	if day < 10:
		day = '0' + str(day)
	report_date = u"{}年{}月{}日".format(now.year,month,day)
	
	context = dict_patient_info
	context['ReportDate'] = report_date
	
##  报_核型
	result = []
	for barcode in report_barcodes:
		tmp_dict = {'sample_id': dict_sample_info.get(barcode, barcode),
					'sample_barcode': barcode,
					'test_result' : dict_result[barcode]
		}
		result.append(tmp_dict)
	context['result'] = result

##  报_图片
#   确定报告中用的红蓝图
	dict_report_png = {}
	report_png = ''
	for barcode in result_barcodes:
		red_blue_no_XY = dict_red_blue_png_no_XY[barcode]
		red_blue_XY = dict_red_blue_png_XY[barcode]
		if dict_result_gender[barcode] == "XX" or dict_result_gender[barcode] == "XY":
			report_png = red_blue_no_XY
		elif dict_result[barcode] == "N/A":
			report_png = red_blue_no_XY
		else:
			report_png = red_blue_XY

		dict_report_png[barcode] = report_png

	subdoc_picture = ReportTML.new_subdoc()
	for barcode in report_barcodes:
		doc_png = dict_report_png[barcode]
		if os.path.exists(doc_png):
			subdoc_picture.add_picture(doc_png,width=Inches(6.1))
		else:
			log = "save report file: " + doc_png + " does not exist!"
			logger.warning(log)	
	context['subdoc_picture'] = subdoc_picture

#   写入，报告完成
	ReportTML.render(context)
	ReportTML.save(out_report)
	logger.info("DONE:save report file")

if __name__ == '__main__':
    main()


# 另一种报告放图片的判断逻辑
'''
if dict_result[barcode] == u"N/A":
	report_png = png_fullpath_no_XY
elif u"多条染色体异常" in dict_result[barcode]:
	if 'X' in dict_result[barcode] or 'Y' in dict_result[barcode]:
		report_png = red_blue_XY
	else:
		report_png = red_blue_no_XY
else:
	report_png = red_blue_XY
'''
